import gradio as gr
from utils.document_processor import ADGMDocumentProcessor as DocumentProcessor
from utils.rag_handler import RAGHandler
from utils.checklist_verifier import ChecklistVerifier
import os
import json
from docx import Document
from typing import List, Dict, Optional
from pathlib import Path


from docx.text.paragraph import Paragraph

def add_comment(self, comment_text: str, author: str = "Reviewer"):
    """
    Simulate adding a comment by appending it inside the paragraph text.
    This keeps your code working without changing your logic.
    """
    self.text += f"  [COMMENT by {author}: {comment_text}]"

Paragraph.add_comment = add_comment

class ADGMComplianceAgent:
    def __init__(self):
        self.doc_processor = DocumentProcessor()
        self.rag_handler = RAGHandler()
        self.checklist_verifier = ChecklistVerifier()
        
        
        try:
            docs_path = "data/adgm_documents"
            if not Path(docs_path).exists():
                raise FileNotFoundError(f"ADGM documents directory not found at {docs_path}")
                
            self.rag_handler.initialize_vectorstore(docs_path)
            self.rag_handler.initialize_qa_chain()
        except Exception as e:
            raise RuntimeError(f"Failed to initialize RAG system: {str(e)}")

    def process_document(self, file_path: str) -> Dict:
        """Process a single document with validation"""
        try:
            doc = Document(file_path)
            doc_text = "\n".join(p.text for p in doc.paragraphs if p.text)
            
            if not doc_text.strip():
                return {"error": "Empty document", "file": file_path}
                
            doc_type = self.doc_processor.identify_document_type(doc_text)
            analysis = self.rag_handler.query(
                f"Analyze this {doc_type} document for ADGM compliance: {doc_text[:2000]}"
            )
            
            return {
                "document": doc_type,
                "analysis": analysis.get("result", "Analysis failed"),
                "file_path": file_path
            }
        except Exception as e:
            return {"error": str(e), "file": file_path}

    def generate_output(self, files: List[str]) -> Dict:
        """Process all files and generate compliance report"""
        results = []
        uploaded_docs = []
        
        for file in files:
            file_path = file.name if hasattr(file, 'name') else file
            result = self.process_document(file_path)
            
            if "error" not in result:
                uploaded_docs.append({
                    "path": file_path,
                    "type": result["document"]
                })
            results.append(result)
        
        
        process_type = self.detect_process_type(uploaded_docs)
        
        
        checklist_result = self.checklist_verifier.verify_documents(
            process_type, 
            uploaded_docs
        )
        
        
        marked_files = []
        for result in results:
            if "file_path" in result:
                try:
                    marked_path = self.create_marked_copy(
                        result["file_path"],
                        result.get("analysis", "")
                    )
                    marked_files.append(marked_path)
                except Exception as e:
                    results.append({
                        "error": f"Markup failed: {str(e)}",
                        "file": result["file_path"]
                    })
        
        return {
            "document_analysis": results,
            "checklist_verification": checklist_result,
            "marked_files": marked_files
        }

    def detect_process_type(self, docs: List[Dict]) -> str:
        """Detect ADGM process type from documents"""
        doc_types = {doc["type"] for doc in docs}
        
        if "Articles of Association" in doc_types:
            return "Company Incorporation"
        elif "Employment Contract" in doc_types:
            return "Employment Compliance"
        return "General Compliance"

    def create_marked_copy(self, original_path: str, comment: str) -> str:
        """Create reviewed document with comments"""
        marked_path = str(Path(original_path).with_stem(
            f"{Path(original_path).stem}_reviewed"
        ))
        
        doc = Document(original_path)
        if doc.paragraphs:
            doc.paragraphs[0].add_comment(
                f"ADGM Compliance Review:\n{comment[:1000]}",
                author="ADGM Agent"
            )
        doc.save(marked_path)
        return marked_path

def launch_interface():
    agent = ADGMComplianceAgent()
    
    def wrapped_generate(files):
        result = agent.generate_output(files)
        return (
            json.dumps(result, indent=2),
            result["marked_files"][0] if result["marked_files"] else None
        )
    
    return gr.Interface(
        fn=wrapped_generate,
        inputs=gr.File(
            file_count="multiple",
            file_types=[".docx"],
            label="Upload ADGM Documents"
        ),
        outputs=[
            gr.JSON(label="Compliance Report"),
            gr.File(label="Reviewed Document")
        ],
        title="ADGM Corporate Agent",
        description="Upload legal documents for ADGM compliance review",
        allow_flagging="never"
    )

if __name__ == "__main__":
    interface = launch_interface()
    interface.launch(server_port=7860)
