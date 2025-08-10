import re
from typing import Dict, List, Optional
from docx import Document
from docx.document import Document as DocxDocument
from pathlib import Path

class ADGMDocumentProcessor:
    """Lightweight document processor with fail-safes"""
    
    DOC_TYPES = {
        'articles': 'Articles of Association',
        'memorandum': 'Memorandum of Association',
        'ubo': 'UBO Declaration',
        'resolution': 'Board Resolution',
        'application': 'Incorporation Application'
    }

    @staticmethod
    def identify_doc_type(text: str) -> Optional[str]:
        """Safer document type detection using regex"""
        text = text.lower().replace(" ", "")
        for pattern, doc_type in ADGMDocumentProcessor.DOC_TYPES.items():
            if re.search(rf'\b{pattern}\b', text):
                return doc_type
        return None

    @staticmethod
    def safe_read_docx(path: str) -> Optional[DocxDocument]:
        """Protected document reading"""
        try:
            if Path(path).suffix.lower() != '.docx':
                raise ValueError("Not a .docx file")
            return Document(path)
        except Exception as e:
            print(f"Error reading {path}: {str(e)}")
            return None

    @staticmethod
    def analyze_document(path: str) -> Dict[str, List[str]]:
        """Extract document structure with validation"""
        doc =ADGMDocumentProcessor.safe_read_docx(path)
        if not doc:
            return {"error": ["Invalid document"]}
            
        structure = {}
        current_section = "Header"
        
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    current_section = para.text
                structure.setdefault(current_section, []).append(para.text)
        
        return structure

    @staticmethod
    def add_review_comment(path: str, section: str, comment: str) -> bool:
        """Non-destructive comment adding"""
        try:
            doc = Document(path)
            for para in doc.paragraphs:
                if para.text.strip() == section:
                    para.add_comment(comment[:500], author="ADGM-AI")
                    doc.save(f"{Path(path).stem}_reviewed.docx")
                    return True
            return False
        except Exception as e:
            print(f"Comment failed: {str(e)}")
            return False